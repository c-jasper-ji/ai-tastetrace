from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
BUILD = ROOT / "docs" / ".pdf-build"
DOCX_PATH = BUILD / "TASTETRACE_PROJECT_GUIDE.docx"

INK = "EAF0F6"
NAVY = "0A0E14"
SURFACE = "111923"
LINE = "2A3746"
MUTED = "8090A2"
VIOLET = "8B62FF"
LIME = "C7F22C"
CYAN = "55D8F4"
WHITE = "FFFFFF"
BLACK = "111827"
LIGHT = "F4F7FA"
LIGHT_LINE = "D8E0EA"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, xy, fill: str, outline: str, radius: int = 18, width: int = 2):
    draw.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def arrow(draw: ImageDraw.ImageDraw, start, end, color: str = CYAN, width: int = 5):
    draw.line([start, end], fill=f"#{color}", width=width)
    x, y = end
    draw.polygon([(x, y), (x - 15, y - 9), (x - 15, y + 9)], fill=f"#{color}")


def make_flow_diagram(path: Path):
    image = Image.new("RGB", (1500, 420), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    title = font(34, True)
    label = font(24, True)
    small = font(18)
    draw.text((50, 32), "TasteTrace user sequence", font=title, fill=f"#{WHITE}")
    boxes = [
        ("1", "Seed tracks", "3-10 listening anchors"),
        ("2", "Discovery brief", "Mood, era, niche intent"),
        ("3", "Hybrid score", "Six visible signals"),
        ("4", "Diversity rerank", "Artist and album controls"),
        ("5", "Evidence output", "Why, sources, feedback"),
    ]
    x = 50
    for index, (number, heading, detail) in enumerate(boxes):
        width = 245
        rounded_box(draw, (x, 130, x + width, 320), SURFACE, LINE)
        draw.ellipse((x + 18, 150, x + 62, 194), fill=f"#{VIOLET}")
        draw.text((x + 32, 157), number, font=font(21, True), fill=f"#{WHITE}", anchor="mm")
        draw.text((x + 18, 220), heading, font=label, fill=f"#{WHITE}")
        draw.multiline_text((x + 18, 260), detail, font=small, fill=f"#{MUTED}", spacing=5)
        if index < len(boxes) - 1:
            arrow(draw, (x + width + 8, 225), (x + width + 43, 225))
        x += 290
    image.save(path)


def make_architecture_diagram(path: Path):
    image = Image.new("RGB", (1500, 760), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 30), "Scalable system architecture", font=font(34, True), fill=f"#{WHITE}")
    nodes = [
        (60, 130, 300, 250, "React UI", "Seeds, controls, evidence"),
        (400, 130, 680, 250, "Recommendation API", "Typed orchestration"),
        (790, 70, 1100, 190, "Identity service", "MusicBrainz IDs"),
        (790, 230, 1100, 350, "Candidate service", "Last.fm / ListenBrainz"),
        (790, 390, 1100, 510, "Feature store", "Tags, embeddings, provenance"),
        (1190, 230, 1440, 350, "Review evidence", "Permissioned adapters"),
        (400, 420, 680, 540, "Hybrid ranker", "Relevance + feedback"),
        (400, 590, 680, 710, "Diversity reranker", "MMR + hard caps"),
        (60, 500, 300, 620, "Feedback store", "Private user signals"),
    ]
    for x1, y1, x2, y2, heading, detail in nodes:
        rounded_box(draw, (x1, y1, x2, y2), SURFACE, LINE)
        draw.text((x1 + 18, y1 + 24), heading, font=font(22, True), fill=f"#{WHITE}")
        draw.text((x1 + 18, y1 + 66), detail, font=font(17), fill=f"#{MUTED}")
    arrow(draw, (300, 190), (400, 190))
    arrow(draw, (680, 190), (790, 130))
    arrow(draw, (680, 200), (790, 290))
    arrow(draw, (680, 210), (790, 450))
    arrow(draw, (1100, 290), (1190, 290))
    arrow(draw, (540, 250), (540, 420), VIOLET)
    arrow(draw, (540, 540), (540, 590), LIME)
    arrow(draw, (300, 560), (400, 490), CYAN)
    image.save(path)


def make_ranking_diagram(path: Path):
    image = Image.new("RGB", (1500, 520), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 30), "Recommendation and diversity pipeline", font=font(34, True), fill=f"#{WHITE}")
    columns = [
        ("Candidate", ["Sonic fit", "Tag overlap", "Scene proximity"]),
        ("Personal", ["Natural-language brief", "Mood and era", "Feedback"]),
        ("Evidence", ["Curated index", "Momentum", "Long-tail value"]),
        ("Set quality", ["Similarity to selected", "Artist cap", "Album penalty"]),
    ]
    x = 50
    colors = [VIOLET, CYAN, LIME, "FFB44A"]
    for i, (heading, rows) in enumerate(columns):
        rounded_box(draw, (x, 120, x + 310, 430), SURFACE, LINE)
        draw.rectangle((x, 120, x + 310, 175), fill=f"#{colors[i]}")
        draw.text((x + 18, 134), heading, font=font(23, True), fill=f"#{NAVY if i in (1, 2, 3) else WHITE}")
        y = 215
        for row in rows:
            draw.ellipse((x + 20, y + 6, x + 32, y + 18), fill=f"#{colors[i]}")
            draw.text((x + 48, y), row, font=font(19), fill=f"#{INK}")
            y += 62
        if i < len(columns) - 1:
            arrow(draw, (x + 318, 275), (x + 355, 275))
        x += 370
    image.save(path)


def make_build_sequence(path: Path):
    image = Image.new("RGB", (1500, 920), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 28), "Build and release sequence", font=font(34, True), fill=f"#{WHITE}")
    phases = [
        ("01", "Problem and evidence", "Define user pain, claims and success measures."),
        ("02", "Catalogue and identity", "Curate tracks and prepare canonical IDs."),
        ("03", "Hybrid recommender", "Implement ranking and grounded explanations."),
        ("04", "Diversity layer", "Add MMR, artist caps and emerging priority."),
        ("05", "Review context", "Add permissioned sources or outbound links."),
        ("06", "Evaluation", "Test relevance, novelty, coverage and exposure."),
        ("07", "Public release", "Audit, license, GitHub Pages and demo."),
    ]
    y = 120
    for index, (number, heading, detail) in enumerate(phases):
        color = [VIOLET, CYAN, LIME, "FFB44A", VIOLET, CYAN, LIME][index]
        rounded_box(draw, (120, y, 1380, y + 88), SURFACE, LINE, 14)
        draw.ellipse((55, y + 16, 111, y + 72), fill=f"#{color}")
        draw.text((83, y + 44), number, font=font(20, True), fill=f"#{NAVY}", anchor="mm")
        draw.text((155, y + 14), heading, font=font(23, True), fill=f"#{WHITE}")
        draw.text((155, y + 50), detail, font=font(18), fill=f"#{MUTED}")
        if index < len(phases) - 1:
            draw.line((83, y + 72, 83, y + 120), fill=f"#{LINE}", width=4)
        y += 112
    image.save(path)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LIGHT_LINE, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=BLACK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def add_text(doc, text, size=10.5, color=BLACK, bold=False, italic=False, after=6, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.18
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_kicker(doc, text):
    return add_text(doc, text.upper(), size=9, color=VIOLET, bold=True, after=4)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        cell.width = Inches(widths[i])
        shade_cell(cell, SURFACE)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        set_run(run, size=9, color=WHITE, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cell = cells[i]
            cell.width = Inches(widths[i])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(str(value))
            set_run(run, size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, accent=VIOLET):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.35)
    shade_cell(cell, "F6F3FF" if accent == VIOLET else "F2FAFC")
    set_cell_border(cell, accent, "8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(label.upper())
    set_run(run, size=8.5, color=accent, bold=True)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    run = paragraph.add_run(text)
    set_run(run, size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("TasteTrace | Project Guide | ")
    set_run(run, size=8, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    heading_tokens = {
        "Heading 1": (16, VIOLET, 14, 7),
        "Heading 2": (13, "344A63", 11, 5),
        "Heading 3": (11, "344A63", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "TASTETRACE  /  EXPLAINABLE MUSIC DISCOVERY"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.runs[0], size=8, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])


def cover(doc: Document):
    for _ in range(5):
        add_text(doc, "", after=7)
    add_kicker(doc, "AI product prototype · Project specification")
    title = add_text(doc, "TasteTrace", size=34, color=NAVY, bold=True, after=2)
    title.paragraph_format.keep_with_next = True
    add_text(
        doc,
        "Explainable, diversity-aware music discovery",
        size=18,
        color="344A63",
        bold=True,
        after=18,
    )
    add_text(
        doc,
        "A functional web application that combines hybrid recommendation, long-tail artist discovery, diversity reranking, and traceable review evidence.",
        size=12,
        color=MUTED,
        after=28,
    )
    add_callout(
        doc,
        "Core proposition",
        "Help listeners find music they are likely to value without repeating the same artists, over-rewarding popularity, or hiding the recommendation logic.",
        VIOLET,
    )
    add_text(doc, "", after=44)
    add_text(doc, "Prepared for Imperial College London", size=11, color=NAVY, bold=True, after=3)
    add_text(doc, "Big Data, AI & Machine Learning · Group Project", size=10, color=MUTED, after=3)
    add_text(doc, "Version 1.0 · June 2026", size=10, color=MUTED, after=0)
    doc.add_page_break()


def build_document():
    BUILD.mkdir(parents=True, exist_ok=True)
    flow = BUILD / "user-sequence.png"
    architecture = BUILD / "architecture.png"
    ranking = BUILD / "ranking.png"
    sequence = BUILD / "build-sequence.png"
    make_flow_diagram(flow)
    make_architecture_diagram(architecture)
    make_ranking_diagram(ranking)
    make_build_sequence(sequence)

    doc = Document()
    configure_document(doc)
    cover(doc)

    add_kicker(doc, "01 · Executive summary")
    add_heading(doc, "Product and business case", 1)
    add_text(
        doc,
        "TasteTrace is a public-ready prototype for listeners who want more adventurous and accountable music discovery. The product starts with a small set of songs the user already values, converts those songs into a transparent taste representation, ranks a broad catalogue, and reranks the final set to reduce artist and album repetition.",
    )
    add_callout(
        doc,
        "Business problem",
        "Existing recommendation feeds reduce search effort but often create repetition, popularity bias, weak explainability, and limited visibility for growing artists.",
        CYAN,
    )
    add_heading(doc, "How the prototype satisfies the assignment", 2)
    add_table(
        doc,
        ["Requirement", "TasteTrace evidence"],
        [
            ("Real business problem", "Selection overload, recommendation repetition, and weak trust."),
            ("Functional software", "Interactive React application with five working views."),
            ("AI component", "Natural-language preference interpretation and grounded explanations."),
            ("ML component", "Feature similarity, hybrid ranking, feedback, and diversity reranking."),
            ("Data component", "Structured catalogue, evidence metadata, source links, and scalable architecture."),
            ("Product management", "Clear users, scope, risks, success measures, and release sequence."),
        ],
        [1.75, 4.6],
    )
    add_heading(doc, "Prototype status", 2)
    for item in [
        "Credential-free and reproducible.",
        "More than 80 curated tracks and more than 70 distinct artists.",
        "One-artist-per-result-set constraint in the default recommendation set.",
        "Outbound RYM, AOTY, and Pitchfork searches without copied live scores.",
        "Automated tests, production build, and zero known npm audit vulnerabilities.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_kicker(doc, "02 · User problem")
    add_heading(doc, "Why a different discovery experience is useful", 1)
    add_table(
        doc,
        ["Pain point", "Observed product pattern", "TasteTrace response"],
        [
            ("Repetition", "Same artist, album, or scene dominates the set.", "MMR reranking plus hard artist and album penalties."),
            ("Popularity loop", "Safe, familiar catalogue receives disproportionate exposure.", "Long-tail value and emerging-artist priority."),
            ("Black-box ranking", "Users see a result but not the evidence.", "Seven visible component scores and bridge seeds."),
            ("Fragmented context", "Recommendation and critical context live separately.", "Direct review-source searches on every result."),
            ("Flat taste labels", "A user becomes one genre or mood.", "Multi-dimensional taste profile and controlled blind spot."),
        ],
        [1.25, 2.2, 2.9],
    )
    add_heading(doc, "Target users", 2)
    for item in [
        "Curious listeners who want a faster route into unfamiliar music.",
        "Heavy listeners who already understand genres and want deeper long-tail discovery.",
        "Users who value editorial or community evidence before committing listening time.",
        "Music platforms seeking a trust, diversity, or catalogue-coverage feature.",
        "Independent and growing artists who benefit from more balanced exposure.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Value hypothesis", 2)
    add_text(
        doc,
        "A more transparent and diverse recommendation set should increase useful discovery, catalogue coverage, saves, session depth, and trust. These are hypotheses for future evaluation, not claimed results.",
    )
    doc.add_page_break()

    add_kicker(doc, "03 · Product sequence")
    add_heading(doc, "End-to-end user journey", 1)
    doc.add_picture(str(flow), width=Inches(6.55))
    add_text(doc, "Figure 1. The complete interactive recommendation sequence.", size=8.5, color=MUTED, italic=True, after=10)
    add_heading(doc, "Working screens", 2)
    add_table(
        doc,
        ["View", "Primary job"],
        [
            ("Discover", "Select seeds, write a brief, control exploration and inspect ranked tracks."),
            ("Artist Radar", "Turn matching tracks into niche and growing artist discovery."),
            ("Taste Profile", "Explain stable preferences, feature centre, and underexplored directions."),
            ("Evidence", "Trace each result to model evidence and external review searches."),
            ("Model Lab", "Adjust weights and observe relevance-diversity trade-offs."),
        ],
        [1.35, 5.0],
    )
    add_heading(doc, "Feedback loop", 2)
    add_text(
        doc,
        "Like, dislike, and save actions persist in the browser. They immediately adjust the deterministic ranking, demonstrating the interaction contract for a later learning-to-rank system.",
    )
    doc.add_page_break()

    add_kicker(doc, "04 · System architecture")
    add_heading(doc, "Current and production architecture", 1)
    doc.add_picture(str(architecture), width=Inches(6.55))
    add_text(doc, "Figure 2. Target architecture with independent identity, candidate, feature, ranking, evidence, and feedback services.", size=8.5, color=MUTED, italic=True, after=10)
    add_heading(doc, "Current prototype", 2)
    add_text(
        doc,
        "The implemented version runs entirely in the browser. This makes the classroom and public demo reliable, prevents credential leakage, and keeps every recommendation reproducible.",
    )
    add_heading(doc, "Production transition", 2)
    for item in [
        "Use MusicBrainz recording and release-group IDs as canonical keys.",
        "Add cached Last.fm and ListenBrainz candidate adapters.",
        "Move features and provenance into a queryable store.",
        "Version the ranker and reranker independently.",
        "Store feedback only with explicit privacy controls.",
    ]:
        add_number(doc, item)
    doc.add_page_break()

    add_kicker(doc, "05 · Recommendation model")
    add_heading(doc, "Hybrid ranking and diversity reranking", 1)
    doc.add_picture(str(ranking), width=Inches(6.55))
    add_text(doc, "Figure 3. Relevance and set quality are separate stages.", size=8.5, color=MUTED, italic=True, after=8)
    add_heading(doc, "Base formula", 2)
    add_callout(
        doc,
        "Like index",
        "30% sonic fit + 20% tags + 14% scene + 12% brief + 10% evidence + 14% discovery - popularity penalty + feedback adjustment.",
        VIOLET,
    )
    add_heading(doc, "Diversity stage", 2)
    add_text(
        doc,
        "The highest-scoring track is selected first. Each subsequent candidate loses score when it resembles tracks already selected. Same-artist and same-album penalties are stronger than ordinary redundancy, while an emerging-artist bonus improves coverage when the set has none.",
    )
    add_heading(doc, "Interpretation", 2)
    for item in [
        "The 0-100 output is a fit index, not a probability.",
        "Evidence strength is not equivalent to personal relevance.",
        "Momentum is displayed separately from model fit.",
        "A lower-ranked candidate may improve the set by adding a new artist or scene.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_kicker(doc, "06 · Catalogue and evidence")
    add_heading(doc, "Data design and source governance", 1)
    add_table(
        doc,
        ["Data group", "Prototype implementation", "Production replacement"],
        [
            ("Identity", "Curated title, artist, album, year, country.", "MusicBrainz IDs and confidence-aware resolution."),
            ("Tags", "Curated genre, mood, and scene labels.", "Permissioned tags plus multilingual curation."),
            ("Audio style", "Deterministically derived feature values.", "Licensed features or local audio embeddings."),
            ("Popularity", "Prototype long-tail feature.", "Dated, source-specific popularity measures."),
            ("Momentum", "Prototype emerging-artist feature.", "Time-windowed growth with provenance."),
            ("Reviews", "Outbound RYM, AOTY, Pitchfork searches.", "Licensed adapters with attribution and timestamps."),
        ],
        [1.25, 2.35, 2.75],
    )
    add_heading(doc, "Why live review scores are not copied", 2)
    add_text(
        doc,
        "Review websites have different score definitions, access rules, and update schedules. Copying visible values without a stable authorised interface would create legal, technical, and provenance risks. The prototype therefore keeps its own clearly labelled evidence index separate and directs users to the original source.",
    )
    add_heading(doc, "Big Data position", 2)
    add_text(
        doc,
        "The bundled catalogue is not presented as big data. The project demonstrates how high-volume metadata, listening events, embeddings, reviews, and feedback would be joined and governed in a production pipeline.",
    )
    doc.add_page_break()

    add_kicker(doc, "07 · Product differentiation")
    add_heading(doc, "What is meaningfully different", 1)
    add_table(
        doc,
        ["Capability", "Typical feed", "TasteTrace"],
        [
            ("Recommendation logic", "Hidden or minimally explained.", "Component scores and bridge evidence."),
            ("Diversity", "Often implicit.", "Visible control and explicit reranking."),
            ("Artist discovery", "Track-first and popularity-heavy.", "Artist Radar with fit, tier, and momentum."),
            ("Review context", "Separate browsing task.", "Direct source trail for each album."),
            ("Taste reflection", "Genre summaries.", "Feature centre plus controlled blind spot."),
            ("Model experimentation", "Not user-accessible.", "Interactive Model Lab."),
        ],
        [1.35, 2.1, 2.9],
    )
    add_heading(doc, "Design language", 2)
    add_text(
        doc,
        "The interface adopts a dark research-console aesthetic. Violet indicates model state, lime indicates discovery, cyan indicates evidence, and semantic colours are reserved for feedback. Dense information hierarchy supports serious comparison without becoming a dashboard of decorative cards.",
    )
    add_heading(doc, "Accessibility and responsiveness", 2)
    for item in [
        "Keyboard focus states on interactive controls.",
        "Text labels for icon-only actions.",
        "Responsive navigation and stacked mobile layouts.",
        "No dependence on album artwork for identity or meaning.",
        "Colour is supported by labels and numeric values.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_kicker(doc, "08 · Evaluation")
    add_heading(doc, "How the product should be tested", 1)
    add_table(
        doc,
        ["Evaluation layer", "Metrics or checks"],
        [
            ("Functional", "Search, seed selection, settings, feedback, navigation, review links."),
            ("Relevance", "Precision@k, recall@k, NDCG, saves, rejection rate."),
            ("Diversity", "Intra-list diversity, unique artists, unique albums, catalogue coverage."),
            ("Discovery", "Novelty, popularity distribution, emerging-artist exposure."),
            ("Trust", "Explanation usefulness, correction behaviour, source-link use."),
            ("Reliability", "Tests, build, audit, API fallback, deterministic fixture."),
            ("Fairness", "Exposure by region, language, audience tier, and catalogue segment."),
        ],
        [1.55, 4.8],
    )
    add_heading(doc, "Current verification", 2)
    for item in [
        "Four automated recommendation-engine tests pass.",
        "TypeScript production build passes.",
        "npm dependency audit reports zero known vulnerabilities.",
        "Final sets contain distinct artists under default constraints.",
        "Scores remain bounded and review links are generated.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "No unsupported results",
        "No user study, online experiment, or production accuracy result has been conducted. The documentation therefore presents evaluation plans rather than invented outcomes.",
        CYAN,
    )
    doc.add_page_break()

    add_kicker(doc, "09 · Responsible AI")
    add_heading(doc, "Privacy, representation, and artist exposure", 1)
    add_heading(doc, "Preference privacy", 2)
    add_text(
        doc,
        "Music preference can correlate with culture, language, mood, religion, age, and community identity. The prototype stores only local feedback. A production service must provide consent, deletion, export, retention controls, and a prohibition on sensitive-trait inference.",
    )
    add_heading(doc, "Catalogue representation", 2)
    add_text(
        doc,
        "The expanded catalogue covers more countries and scenes than the original trial, but it remains incomplete and curator-dependent. Audience tiers and tags are subjective. Future work requires multilingual curation and exposure audits.",
    )
    add_heading(doc, "Artist exposure", 2)
    add_text(
        doc,
        "Long-tail bonuses and diversity constraints can improve exposure, but they can also be manipulated or create arbitrary thresholds. Momentum must be time-windowed, source-dated, and separated from personal fit.",
    )
    add_heading(doc, "Human agency", 2)
    for item in [
        "Users can inspect why a result appears.",
        "Users can change exploration and diversity.",
        "Users can reject or save results.",
        "Users can inspect independent external sources.",
        "The system avoids objective-quality language.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_kicker(doc, "10 · Build sequence")
    add_heading(doc, "From prototype to open-source product", 1)
    doc.add_picture(str(sequence), width=Inches(6.55))
    add_text(doc, "Figure 4. Controlled implementation and release order.", size=8.5, color=MUTED, italic=True, after=8)
    add_text(
        doc,
        "The prompt pack in the repository mirrors this sequence. Each phase has an objective, constraints, verification, and completion condition so future AI-assisted coding remains auditable.",
    )
    doc.add_page_break()

    add_kicker(doc, "11 · Delivery and roadmap")
    add_heading(doc, "Current deliverables", 1)
    for item in [
        "Working responsive web application.",
        "Expanded curated catalogue.",
        "Hybrid recommendation and diversity model.",
        "Artist Radar and review evidence workflow.",
        "English README and technical documentation.",
        "Structured AI coding prompt pack.",
        "Automated tests and production build.",
        "This visual project specification.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Next release phase", 2)
    add_number(doc, "Choose an open-source licence.")
    add_number(doc, "Initialise Git and review generated artefacts.")
    add_number(doc, "Add GitHub repository metadata and contribution guidance.")
    add_number(doc, "Configure GitHub Actions for tests and build.")
    add_number(doc, "Publish the static build through GitHub Pages.")
    add_number(doc, "Run the documented demo against the public URL.")
    add_heading(doc, "Presentation order", 2)
    add_text(
        doc,
        "Problem -> seed and brief -> recommendation evidence -> diversity constraint -> Artist Radar -> review trail -> Taste Profile -> Model Lab -> limitations -> roadmap.",
    )
    add_callout(
        doc,
        "Final message",
        "TasteTrace does not try to replace human taste or criticism. It makes machine discovery more inspectable, more diverse, and easier to challenge.",
        VIOLET,
    )
    doc.add_page_break()

    add_kicker(doc, "Appendix")
    add_heading(doc, "Key source and extension references", 1)
    sources = [
        ("MusicBrainz", "https://musicbrainz.org/doc/MusicBrainz_API"),
        ("ListenBrainz", "https://listenbrainz.readthedocs.io/"),
        ("Last.fm API", "https://www.last.fm/api"),
        ("Rate Your Music", "https://rateyourmusic.com/"),
        ("Album of the Year", "https://www.albumoftheyear.org/"),
        ("Pitchfork", "https://pitchfork.com/"),
        ("Maximal marginal relevance", "Carbonell and Goldstein (1998), diversification of information retrieval."),
    ]
    add_table(doc, ["Reference", "Use"], sources, [2.0, 4.35])
    add_heading(doc, "Repository commands", 2)
    add_callout(
        doc,
        "Local verification",
        "cd app  |  npm install  |  npm test  |  npm run build  |  npm audit",
        CYAN,
    )
    add_text(
        doc,
        "Prepared from the implemented repository state. Public GitHub publication is intentionally reserved for the next project conversation.",
        size=9,
        color=MUTED,
        italic=True,
        after=0,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()

