from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
BUILD = ROOT / "docs" / ".model-pdf-build"
DOCX_PATH = BUILD / "TASTETRACE_MODEL_TECHNICAL_SPECIFICATION.docx"

NAVY = "0A0E14"
SURFACE = "111923"
LINE = "2A3746"
MUTED = "6F7F91"
INK = "17202A"
WHITE = "FFFFFF"
VIOLET = "7653E8"
CYAN = "2DAFCB"
LIME = "A6CB16"
ORANGE = "E79A35"
LIGHT = "F5F7FA"
LIGHT_LINE = "D8E0EA"
PALE_VIOLET = "F2EEFF"
PALE_CYAN = "EAF8FB"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def image_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, xy, fill: str, outline: str, radius=18, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def arrow(draw: ImageDraw.ImageDraw, start, end, color=CYAN, width=5):
    draw.line([start, end], fill=f"#{color}", width=width)
    x, y = end
    draw.polygon([(x, y), (x - 16, y - 9), (x - 16, y + 9)], fill=f"#{color}")


def architecture_diagram(path: Path):
    image = Image.new("RGB", (1500, 730), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 28), "Four-layer recommendation architecture", font=image_font(36, True), fill=f"#{WHITE}")
    layers = [
        (70, 125, 350, 605, "1", "IDENTITY", "MusicBrainz", VIOLET,
         ["Recording MBID", "Artist and release", "Tags and genres", "Relationships"]),
        (420, 125, 700, 605, "2", "FEATURES", "FMA + Essentia", CYAN,
         ["Audio embeddings", "Tempo and tonality", "Mood probabilities", "Versioned features"]),
        (770, 125, 1050, 605, "3", "RETRIEVAL", "TasteTrace KNN", LIME,
         ["Query centroid", "Mixed similarity", "Top-k candidates", "ANN at scale"]),
        (1120, 125, 1400, 605, "4", "RERANKING", "TasteTrace MMR", ORANGE,
         ["Relevance-diversity", "Artist cap", "Album penalty", "Final evidence"]),
    ]
    for index, (x1, y1, x2, y2, number, label, heading, color, rows) in enumerate(layers):
        rounded_box(draw, (x1, y1, x2, y2), SURFACE, LINE)
        draw.ellipse((x1 + 20, y1 + 20, x1 + 72, y1 + 72), fill=f"#{color}")
        draw.text((x1 + 46, y1 + 46), number, font=image_font(22, True), fill=f"#{NAVY}", anchor="mm")
        draw.text((x1 + 20, y1 + 100), label, font=image_font(17, True), fill=f"#{color}")
        draw.text((x1 + 20, y1 + 142), heading, font=image_font(27, True), fill=f"#{WHITE}")
        y = y1 + 225
        for row in rows:
            draw.ellipse((x1 + 22, y + 7, x1 + 34, y + 19), fill=f"#{color}")
            draw.text((x1 + 50, y), row, font=image_font(19), fill="#C8D2DE")
            y += 62
        if index < len(layers) - 1:
            arrow(draw, (x2 + 8, 365), (x2 + 63, 365), color)
    draw.text(
        (70, 660),
        "AcousticBrainz is an optional historical enrichment source, not a live production dependency.",
        font=image_font(19),
        fill="#91A1B2",
    )
    image.save(path)


def feature_diagram(path: Path):
    image = Image.new("RGB", (1500, 700), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 28), "Observation, independent variables and target", font=image_font(36, True), fill=f"#{WHITE}")
    rounded_box(draw, (70, 115, 430, 610), SURFACE, LINE)
    draw.text((95, 145), "Observation i", font=image_font(25, True), fill=f"#{VIOLET}")
    draw.text((95, 200), "One recording", font=image_font(28, True), fill=f"#{WHITE}")
    draw.text((95, 250), "Canonical key:", font=image_font(18), fill="#91A1B2")
    draw.text((95, 282), "MusicBrainz recording MBID", font=image_font(19, True), fill="#D8E1EA")
    draw.text((95, 360), "Seed query q", font=image_font(22, True), fill=f"#{CYAN}")
    draw.multiline_text(
        (95, 405),
        "Weighted centroid of\nselected seed vectors",
        font=image_font(19),
        fill="#C8D2DE",
        spacing=9,
    )
    arrow(draw, (430, 360), (500, 360), CYAN)
    rounded_box(draw, (500, 115, 1015, 610), SURFACE, LINE)
    draw.text((525, 145), "Independent variables X_i", font=image_font(25, True), fill=f"#{LIME}")
    blocks = [
        ("Audio", "embedding, tempo, key, loudness"),
        ("Semantics", "genre, mood and tag vectors"),
        ("Context", "year, country and relationships"),
        ("Quality", "missingness and source confidence"),
    ]
    y = 220
    colors = [VIOLET, CYAN, LIME, ORANGE]
    for (heading, detail), color in zip(blocks, colors):
        draw.rectangle((525, y, 545, y + 58), fill=f"#{color}")
        draw.text((568, y), heading, font=image_font(21, True), fill=f"#{WHITE}")
        draw.text((568, y + 31), detail, font=image_font(17), fill="#AAB7C5")
        y += 92
    arrow(draw, (1015, 360), (1085, 360), ORANGE)
    rounded_box(draw, (1085, 115, 1430, 610), SURFACE, LINE)
    draw.text((1110, 145), "Dependent variable y", font=image_font(24, True), fill=f"#{ORANGE}")
    draw.text((1110, 210), "KNN retrieval", font=image_font(20, True), fill=f"#{WHITE}")
    draw.multiline_text(
        (1110, 250),
        "No required y.\nDistance ranks neighbours.",
        font=image_font(18),
        fill="#C8D2DE",
        spacing=8,
    )
    draw.text((1110, 355), "Optional supervised y(q,i)", font=image_font(20, True), fill=f"#{WHITE}")
    draw.multiline_text(
        (1110, 400),
        "0 = irrelevant\n1 = plausible\n2 = strongly relevant",
        font=image_font(18),
        fill="#C8D2DE",
        spacing=8,
    )
    image.save(path)


def retrieval_diagram(path: Path):
    image = Image.new("RGB", (1500, 700), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 28), "KNN candidate retrieval and MMR list construction", font=image_font(36, True), fill=f"#{WHITE}")
    stages = [
        (55, 135, 280, 530, "Seed set", ["Resolve MBIDs", "Load vectors", "Apply seed weights"], VIOLET),
        (355, 135, 580, 530, "Query vector", ["Weighted centroid", "Standardised blocks", "Missingness flags"], CYAN),
        (655, 135, 880, 530, "KNN / ANN", ["Mixed similarity", "k = 100", "Candidate pool"], LIME),
        (955, 135, 1180, 530, "MMR", ["lambda = 0.65", "Redundancy penalty", "Artist and album caps"], ORANGE),
        (1255, 135, 1470, 530, "Output", ["12 recommendations", "Artist Radar", "Evidence trail"], VIOLET),
    ]
    for index, (x1, y1, x2, y2, heading, rows, color) in enumerate(stages):
        rounded_box(draw, (x1, y1, x2, y2), SURFACE, LINE)
        draw.rectangle((x1, y1, x2, y1 + 68), fill=f"#{color}")
        draw.text((x1 + 18, y1 + 19), heading, font=image_font(23, True), fill=f"#{NAVY if color != VIOLET else WHITE}")
        y = y1 + 125
        for row in rows:
            draw.ellipse((x1 + 18, y + 5, x1 + 30, y + 17), fill=f"#{color}")
            draw.text((x1 + 45, y), row, font=image_font(18), fill="#C8D2DE")
            y += 70
        if index < len(stages) - 1:
            arrow(draw, (x2 + 8, 330), (x2 + 65, 330), color)
    draw.text(
        (55, 595),
        "The retrieval stage optimises individual similarity. The reranking stage optimises the quality of the set.",
        font=image_font(21, True),
        fill="#D7E0E9",
    )
    image.save(path)


def split_diagram(path: Path):
    image = Image.new("RGB", (1500, 510), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 28), "Artist-grouped offline evaluation", font=image_font(36, True), fill=f"#{WHITE}")
    blocks = [
        (70, 140, 900, 330, "TRAINING 70%", "Fit scaler, tag vocabulary, PCA or metric model", VIOLET),
        (900, 140, 1165, 330, "VALIDATION 15%", "Tune k, weights, metric and MMR", CYAN),
        (1165, 140, 1430, 330, "TEST 15%", "Final locked report", LIME),
    ]
    for x1, y1, x2, y2, heading, detail, color in blocks:
        draw.rectangle((x1, y1, x2, y2), fill=f"#{color}", outline=f"#{LINE}", width=2)
        draw.text((x1 + 20, y1 + 30), heading, font=image_font(23, True), fill=f"#{NAVY if color != VIOLET else WHITE}")
        max_width = x2 - x1 - 40
        words = detail.split()
        lines, line = [], ""
        for word in words:
            candidate = f"{line} {word}".strip()
            if draw.textlength(candidate, font=image_font(18)) <= max_width:
                line = candidate
            else:
                lines.append(line)
                line = word
        if line:
            lines.append(line)
        draw.multiline_text((x1 + 20, y1 + 92), "\n".join(lines), font=image_font(18), fill=f"#{NAVY}", spacing=5)
    draw.text(
        (70, 385),
        "Split by artist rather than random track to reduce identity leakage between train and test data.",
        font=image_font(20),
        fill="#AAB7C5",
    )
    image.save(path)


def backend_diagram(path: Path):
    image = Image.new("RGB", (1500, 790), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    draw.text((50, 28), "Production backend and batch lifecycle", font=image_font(36, True), fill=f"#{WHITE}")
    nodes = [
        (55, 125, 300, 245, "React UI", "Public seed and filter input", VIOLET),
        (380, 125, 680, 245, "FastAPI gateway", "Validation and orchestration", CYAN),
        (770, 80, 1060, 200, "MusicBrainz adapter", "MBIDs, metadata and cache", VIOLET),
        (770, 250, 1060, 370, "Feature workers", "Essentia / MusiCNN", CYAN),
        (770, 420, 1060, 540, "Catalogue store", "PostgreSQL + provenance", LIME),
        (1160, 250, 1440, 370, "Vector index", "pgvector / FAISS HNSW", LIME),
        (380, 430, 680, 550, "KNN service", "Top-k candidate retrieval", LIME),
        (380, 610, 680, 730, "MMR service", "Diversity and hard caps", ORANGE),
        (55, 610, 300, 730, "Evidence output", "Scores, sources and reasons", VIOLET),
    ]
    for x1, y1, x2, y2, heading, detail, color in nodes:
        rounded_box(draw, (x1, y1, x2, y2), SURFACE, LINE, 14)
        draw.rectangle((x1, y1, x1 + 12, y2), fill=f"#{color}")
        draw.text((x1 + 28, y1 + 22), heading, font=image_font(21, True), fill=f"#{WHITE}")
        draw.text((x1 + 28, y1 + 68), detail, font=image_font(16), fill="#9FADBC")
    arrow(draw, (300, 185), (380, 185), CYAN)
    arrow(draw, (680, 170), (770, 140), VIOLET)
    arrow(draw, (680, 185), (770, 310), CYAN)
    arrow(draw, (680, 200), (770, 480), LIME)
    arrow(draw, (1060, 310), (1160, 310), LIME)
    arrow(draw, (1300, 370), (680, 490), LIME)
    arrow(draw, (530, 550), (530, 610), ORANGE)
    arrow(draw, (380, 670), (300, 670), VIOLET)
    draw.text((1120, 450), "Nightly / weekly batch", font=image_font(19, True), fill=f"#{ORANGE}")
    draw.multiline_text(
        (1120, 490),
        "ingest -> resolve -> extract\n-> validate -> index -> promote",
        font=image_font(17),
        fill="#AAB7C5",
        spacing=7,
    )
    image.save(path)


def set_run(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LIGHT_LINE, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches):
    widths = [round(width * 1440) for width in widths_inches]
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[index])
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_text(doc, text, size=11, color=INK, bold=False, italic=False, after=6, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_kicker(doc, text):
    return add_text(doc, text.upper(), size=8.5, color=VIOLET, bold=True, after=4)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(text), size=10.5)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(text), size=10.5)
    return paragraph


def add_numbered_sequence(doc, items):
    numbering = doc.part.numbering_part.element
    existing_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId"))
    ]
    new_num_id = max(existing_ids, default=0) + 1
    base_num_id = doc.styles["List Number"].element.pPr.numPr.numId.val
    base_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)

    for text in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = new_num_id
        set_run(paragraph.add_run(text), size=10.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    repeat_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        shade_cell(cell, SURFACE)
        set_cell_border(cell)
        set_cell_margins(cell, top=100, bottom=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_run(paragraph.add_run(value), size=8.5, color=WHITE, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            set_cell_border(cell)
            set_cell_margins(cell, top=95, bottom=95)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.12
            set_run(paragraph.add_run(str(value)), size=8.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, label, text, accent=VIOLET):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_VIOLET if accent == VIOLET else PALE_CYAN)
    set_cell_border(cell, accent, "8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    set_run(p1.add_run(label.upper()), size=8.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    set_run(p2.add_run(text), size=10)
    set_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_formula(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, NAVY)
    set_cell_border(cell, LINE, "5")
    set_cell_margins(cell, top=160, start=200, bottom=160, end=200)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(lines), size=10, color=WHITE, bold=True, name="Consolas")
    set_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("TasteTrace | Model Technical Specification | "), size=8, color=MUTED)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, VIOLET, 18, 10),
        "Heading 2": (13, "31475E", 14, 7),
        "Heading 3": (12, "31475E", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "TASTETRACE  /  MODEL AND DATA ARCHITECTURE"
    set_run(header.runs[0], size=8, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])


def add_picture(doc, path: Path, caption: str, width=6.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_text(doc, caption, size=8.5, color=MUTED, italic=True, after=7)


def build_document():
    BUILD.mkdir(parents=True, exist_ok=True)
    architecture = BUILD / "four-layer-architecture.png"
    variables = BUILD / "variables-and-target.png"
    retrieval = BUILD / "knn-mmr-pipeline.png"
    split = BUILD / "artist-grouped-split.png"
    backend = BUILD / "backend-lifecycle.png"
    architecture_diagram(architecture)
    feature_diagram(variables)
    retrieval_diagram(retrieval)
    split_diagram(split)
    backend_diagram(backend)

    doc = Document()
    configure_document(doc)

    for _ in range(5):
        add_text(doc, "", after=8)
    add_kicker(doc, "Technical model specification")
    add_text(doc, "TasteTrace", size=34, color=NAVY, bold=True, after=1)
    add_text(doc, "Public Cold-Start Music and Artist Finder", size=18, color="31475E", bold=True, after=16)
    add_text(
        doc,
        "Data architecture, variables, K-nearest-neighbour retrieval, diversity reranking, training design, evaluation, and production sequence.",
        size=12,
        color=MUTED,
        after=25,
    )
    add_callout(
        doc,
        "Model position",
        "The current website is a deterministic prototype. This specification defines the academically explicit model and backend target without claiming that the full production pipeline has already been trained.",
        VIOLET,
    )
    add_text(doc, "", after=40)
    add_text(doc, "Prepared for Imperial College London", size=11, color=NAVY, bold=True, after=3)
    add_text(doc, "Big Data, AI & Machine Learning | June 2026 | Version 1.0", size=10, color=MUTED, after=0)
    doc.add_page_break()

    add_kicker(doc, "01 | System definition")
    add_heading(doc, "A public cold-start recommender", 1)
    add_text(
        doc,
        "TasteTrace is designed for any visitor, including users with no account and no imported listening history. The user supplies seed tracks or artists plus optional discovery constraints. The model treats those inputs as query evidence and searches a public catalogue for content-near candidates.",
    )
    add_heading(doc, "Prediction task", 2)
    add_table(
        doc,
        ["Element", "Technical definition"],
        [
            ("Input", "Seed recordings or artists, natural-language brief, mood, era, geography, exploration, and diversity."),
            ("Candidate unit", "A MusicBrainz recording MBID joined to metadata and feature vectors."),
            ("Primary output", "A ranked and diversified set of recordings and associated artists."),
            ("Model family", "Content-based KNN or ANN retrieval followed by MMR reranking."),
            ("Cold-start property", "No personal listening history is required; similarity comes from supplied seeds and public data."),
        ],
        [1.45, 5.05],
    )
    add_callout(
        doc,
        "Important distinction",
        "This is not collaborative filtering. The model can later use public playlist or interaction evidence, but the core finder remains operational from content and metadata alone.",
        CYAN,
    )
    doc.add_page_break()

    add_kicker(doc, "02 | Four-layer design")
    add_heading(doc, "Source-to-recommendation architecture", 1)
    add_picture(doc, architecture, "Figure 1. MusicBrainz identity, acoustic features, KNN retrieval, and MMR reranking.")
    add_text(
        doc,
        "MusicBrainz supplies canonical identity and catalogue relationships. FMA supplies reproducible research data. Local Essentia models create production features where lawful audio or previews are available. AcousticBrainz is treated only as an optional historical archive because its ingestion project ended in 2022.",
    )
    doc.add_page_break()

    add_kicker(doc, "03 | Data responsibilities")
    add_heading(doc, "What each source can and cannot claim", 1)
    add_table(
        doc,
        ["Source", "Primary role", "Do not assume"],
        [
            ("MusicBrainz", "MBIDs, artist/recording/release identity, dates, countries, relationships, community tags.", "Complete audio features, calibrated popularity, or preference labels."),
            ("AcousticBrainz", "Archived low- and high-level descriptors linked by recording MBID.", "A current or continuously expanding production API."),
            ("FMA", "Research metadata, genre labels, features, play counts, and licenced audio subsets.", "Full representation of commercial or global music."),
            ("Essentia / MusiCNN", "Local embeddings, rhythm, tonal, spectral, genre, and mood inference.", "Permission to analyse audio that the project does not lawfully access."),
        ],
        [1.25, 2.75, 2.5],
    )
    add_heading(doc, "MusicBrainz ingestion policy", 2)
    for item in [
        "Use recording and release-group MBIDs as stable join keys.",
        "Cache API responses and send an identifiable User-Agent.",
        "Respect public rate limits; use database dumps for large catalogue refreshes.",
        "Store source dates and confidence for community tags and genres.",
        "Keep cover-art retrieval separate through the Cover Art Archive.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_kicker(doc, "04 | Variables")
    add_heading(doc, "Observation, X and optional y", 1)
    add_picture(doc, variables, "Figure 2. The recording is the observation; feature blocks form X; KNN does not require y.")
    add_formula(
        doc,
        "X_i = [audio_embedding, tempo, key, loudness, mood_probs,\n"
        "       tag_vector, year, country, relationships, confidence]\n\n"
        "q = sum(w_s * X_s) / sum(w_s), for seed tracks s",
    )
    add_text(
        doc,
        "For supervised evaluation or future learning-to-rank, y(q,i) is an ordinal relevance label: 0 irrelevant, 1 plausible, and 2 strongly relevant. MusicBrainz metadata alone is not a preference target.",
    )
    doc.add_page_break()

    add_kicker(doc, "05 | KNN retrieval")
    add_heading(doc, "Mixed-feature nearest neighbours", 1)
    add_formula(
        doc,
        "similarity(q,i) =\n"
        "  0.50 cosine(audio_q, audio_i)\n"
        "+ 0.20 cosine(tags_q, tags_i)\n"
        "+ 0.10 tempo_similarity\n"
        "+ 0.08 tonal_similarity\n"
        "+ 0.07 era_similarity\n"
        "+ 0.05 scene_similarity",
    )
    add_text(
        doc,
        "The coefficients are initial product assumptions, not estimated causal effects. Validation chooses their final values. Continuous variables are standardised using training-only statistics; missing variables receive explicit flags.",
    )
    add_heading(doc, "Why KNN fits the product", 2)
    for item in [
        "It directly answers the user question: what lies near these seeds?",
        "It supports explanations through shared features and nearest seed bridges.",
        "It works without user histories or account-level training data.",
        "It scales from exact search to approximate vector indexes.",
        "It can retrieve tracks or aggregate recording vectors into artist representations.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_kicker(doc, "06 | Advanced parameters")
    add_heading(doc, "Initial KNN and vector-index configuration", 1)
    add_table(
        doc,
        ["Parameter", "Initial", "Validation range or alternative"],
        [
            ("Candidate neighbours k", "100", "25, 50, 100, 200"),
            ("Audio metric", "Cosine", "Cosine or Euclidean"),
            ("Tag encoding", "TF-IDF multi-hot", "Binary or TF-IDF"),
            ("Seed aggregation", "Weighted mean", "Mean, medoid, or max similarity"),
            ("Metadata confidence", ">= 0.50", "0.30 to 0.80"),
            ("Seed-artist exclusion", "Enabled", "Enabled or disabled"),
            ("Prototype index", "Exact search", "Brute force"),
            ("Production index", "HNSW", "pgvector or FAISS"),
            ("HNSW M", "16", "8, 16, 32"),
            ("HNSW ef_search", "64", "32, 64, 128"),
        ],
        [2.05, 1.35, 3.1],
    )
    add_callout(
        doc,
        "Scaling rule",
        "Exact KNN is preferred while the catalogue is small because it is easy to audit. Approximate nearest-neighbour indexing becomes appropriate at hundreds of thousands or millions of recordings.",
        CYAN,
    )
    doc.add_page_break()

    add_kicker(doc, "07 | MMR reranking")
    add_heading(doc, "Relevance is not the same as set quality", 1)
    add_picture(doc, retrieval, "Figure 3. KNN constructs the candidate pool; MMR constructs the final recommendation set.")
    add_formula(
        doc,
        "MMR(i) = lambda * relevance(q,i)\n"
        "       - (1-lambda) * max_similarity(i, selected)\n"
        "       - same_artist_penalty - same_album_penalty\n"
        "       + emerging_artist_bonus",
    )
    add_table(
        doc,
        ["Parameter", "Initial value", "Interpretation"],
        [
            ("lambda", "0.65", "Balance toward relevance while retaining diversity."),
            ("Same artist penalty", "0.40", "Strongly suppress repeated artists."),
            ("Same album penalty", "0.28", "Suppress repeated releases."),
            ("Emerging coverage bonus", "0.035", "Small set-level exposure incentive."),
            ("Artist cap", "1", "Maximum selected results per artist."),
            ("Final set", "12", "Default user-facing recommendations."),
        ],
        [2.0, 1.25, 3.25],
    )
    doc.add_page_break()

    add_kicker(doc, "08 | Training design")
    add_heading(doc, "What is trained in an instance-based model", 1)
    add_text(
        doc,
        "KNN stores observations rather than fitting a conventional regression equation. The training process still fits preprocessing, selects dimensions, optionally learns a distance metric, builds the vector index, and chooses hyperparameters.",
    )
    add_picture(doc, split, "Figure 4. Proposed artist-grouped 70/15/15 split.")
    add_heading(doc, "Dataset use", 2)
    add_table(
        doc,
        ["Partition", "Purpose"],
        [
            ("Training", "Fit scaler, vocabulary, PCA or embedding projection, and optional metric-learning model."),
            ("Validation", "Choose k, feature weights, distance, HNSW parameters, MMR lambda, and penalties."),
            ("Test", "Produce the final locked report once; do not tune against this set."),
        ],
        [1.35, 5.15],
    )
    doc.add_page_break()

    add_kicker(doc, "09 | Evaluation")
    add_heading(doc, "Baselines and success metrics", 1)
    add_heading(doc, "Required baselines", 2)
    add_numbered_sequence(doc, [
        "Popularity-only ranking.",
        "Genre or tag Jaccard similarity.",
        "Audio-only KNN.",
        "Metadata-only KNN.",
        "Hybrid KNN without MMR.",
        "Hybrid KNN with MMR.",
    ])
    add_heading(doc, "Evaluation matrix", 2)
    add_table(
        doc,
        ["Objective", "Metrics"],
        [
            ("Relevance", "Precision@K, Recall@K, MAP@K, NDCG@K."),
            ("Diversity", "Intra-list diversity, unique artist ratio, unique album ratio."),
            ("Discovery", "Catalogue coverage, novelty, long-tail and emerging exposure."),
            ("Reliability", "Latency, cache hit rate, index build success, missing-feature rate."),
            ("Trust", "Explanation usefulness and correction behaviour in a controlled study."),
            ("Representation", "Exposure by country, language, genre, and catalogue segment."),
        ],
        [1.45, 5.05],
    )
    add_callout(
        doc,
        "No fabricated accuracy",
        "The project must report these as planned metrics until a labelled dataset and locked test evaluation have actually been run.",
        VIOLET,
    )
    doc.add_page_break()

    add_kicker(doc, "10 | Backend")
    add_heading(doc, "Production services and batch sequence", 1)
    add_picture(doc, backend, "Figure 5. Online recommendation services plus a versioned offline feature and index pipeline.")
    add_heading(doc, "Controlled batch release", 2)
    add_numbered_sequence(doc, [
        "Ingest MusicBrainz metadata or database-dump changes.",
        "Resolve recording and release identities through MBIDs.",
        "Attach FMA, archived AcousticBrainz, or local Essentia features.",
        "Validate schema, missingness, licence, and provenance.",
        "Fit versioned preprocessing and rebuild the vector index.",
        "Run relevance, diversity, representation, and latency checks.",
        "Promote the index only when acceptance thresholds pass; retain the previous version for rollback.",
    ])
    doc.add_page_break()

    add_kicker(doc, "11 | Implementation boundary")
    add_heading(doc, "Current prototype versus target system", 1)
    add_table(
        doc,
        ["Capability", "Current application", "Specified target"],
        [
            ("Catalogue", "83 curated tracks", "MusicBrainz-backed catalogue and cache"),
            ("Audio variables", "Deterministic tag-derived values", "Essentia embeddings and descriptors"),
            ("Retrieval", "Full-catalogue weighted scoring", "KNN or ANN candidate retrieval"),
            ("Training data", "None", "FMA plus labelled similarity tasks"),
            ("Dependent variable", "None", "Optional ordinal relevance label"),
            ("Diversity", "MMR-style deterministic reranking", "Validated MMR and exposure monitoring"),
            ("Backend", "Browser-only", "FastAPI, PostgreSQL, vector index, workers"),
            ("Updates", "Code release", "Versioned batch and rollback pipeline"),
        ],
        [1.35, 2.35, 2.8],
    )
    add_heading(doc, "Technical risks", 2)
    for item in [
        "MusicBrainz tags and genres are community-generated and incomplete.",
        "FMA is a research corpus, not a complete market representation.",
        "AcousticBrainz coverage is historical and frozen.",
        "Audio extraction requires lawful access to audio or previews.",
        "Similarity should never be presented as objective artistic quality.",
        "All model and index versions require source, date, licence, and transformation provenance.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_kicker(doc, "Appendix | References")
    add_heading(doc, "Primary technical sources", 1)
    add_table(
        doc,
        ["Reference", "Project use"],
        [
            ("MusicBrainz API", "Canonical identity, metadata search, relationships, tags, and genres."),
            ("MusicBrainz database downloads", "Bulk catalogue refresh at scale."),
            ("AcousticBrainz", "Optional historical descriptors; discontinued ingestion noted."),
            ("FMA dataset", "Reproducible feature, genre, training, and evaluation corpus."),
            ("Essentia models", "Local embeddings, classification, auto-tagging, and descriptors."),
            ("Carbonell and Goldstein (1998)", "Maximal marginal relevance diversification."),
        ],
        [2.25, 4.25],
    )
    add_heading(doc, "URLs", 2)
    for url in [
        "https://musicbrainz.org/doc/MusicBrainz_API",
        "https://musicbrainz.org/doc/MusicBrainz_Database/Download",
        "https://acousticbrainz.org/",
        "https://github.com/mdeff/fma",
        "https://essentia.upf.edu/models.html",
    ]:
        add_bullet(doc, url)
    add_callout(
        doc,
        "Repository document",
        "The full text specification is also available in docs/MODEL_TECHNICAL_SPECIFICATION.md.",
        CYAN,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
